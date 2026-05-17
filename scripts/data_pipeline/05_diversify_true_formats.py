"""05_diversify_true_formats.py — Diversify the K8s true-corpus file formats.

After Step 02 scrapes the K8s docs as `.txt`, all signal documents are in a
single format.  Real-world enterprise knowledge bases are a mix of formats
(PDF, DOCX, HTML, TXT), and the RAG ingestion pipeline must handle all of
them.  This script rewrites a portion of the `.txt` corpus into PDF, DOCX,
and HTML so the final true-corpus is a realistic multi-format mix.

Default distribution (deterministic, based on sorted-filename index):
    25 % → .pdf   (reportlab)
    25 % → .docx  (python-docx)
    25 % → .html  (string-template, no external dep)
    25 % → .txt   (kept as-is)

For each file converted to a non-txt format, the original .txt is removed
so the corpus does not contain duplicate content under multiple filenames.

This script is **idempotent** in the loose sense: re-running it after the
mix already exists will skip files whose target format already exists.
Use --force to rebuild from scratch (note: it requires the original .txt
files, which may have been deleted — in that case use --reset).

Usage:
    uv run python scripts/data_pipeline/05_diversify_true_formats.py
    uv run python scripts/data_pipeline/05_diversify_true_formats.py --force
    uv run python scripts/data_pipeline/05_diversify_true_formats.py --dry-run
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

import yaml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
CONFIG_PATH = SCRIPT_DIR / "config.yaml"


def load_config() -> dict:
    with open(CONFIG_PATH) as fh:
        return yaml.safe_load(fh)


def log(msg: str, *, indent: int = 0) -> None:
    prefix = "  " * indent
    print(f"[05_diversify] {prefix}{msg}", flush=True)


def human_size(total_bytes: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if total_bytes < 1024:
            return f"{total_bytes:.1f} {unit}"
        total_bytes /= 1024  # type: ignore[assignment]
    return f"{total_bytes:.1f} TB"


# ---------------------------------------------------------------------------
# Format assignment — deterministic by sorted-index modulo 4
# ---------------------------------------------------------------------------
FORMATS = ("pdf", "docx", "html", "txt")


def assign_format(index: int) -> str:
    """Deterministically map a 0-based index → output format.

    Round-robin assignment ensures an even split regardless of corpus size.
    """
    return FORMATS[index % len(FORMATS)]


# ---------------------------------------------------------------------------
# Converters
# ---------------------------------------------------------------------------

def _title_from_path(path: Path) -> str:
    """Derive a human-readable title from the slug-style filename."""
    stem = path.stem.replace("__", " › ").replace("_", " ")
    return stem.strip(" ›").title()


def write_pdf(src_text: str, src_path: Path, dest: Path) -> None:
    """Render *src_text* as a simple multi-page PDF via reportlab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    body_style = styles["BodyText"]

    doc = SimpleDocTemplate(
        str(dest),
        pagesize=LETTER,
        leftMargin=54, rightMargin=54,
        topMargin=54,  bottomMargin=54,
        title=_title_from_path(src_path),
    )

    story: list = [Paragraph(_title_from_path(src_path), title_style), Spacer(1, 18)]

    # Split on blank lines → one Paragraph per block to keep things simple.
    # reportlab interprets some characters as markup, so escape < > &.
    for block in src_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = (
            block.replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;")
                 .replace("\n", "<br/>")
        )
        story.append(Paragraph(safe, body_style))
        story.append(Spacer(1, 10))

    doc.build(story)


def write_docx(src_text: str, src_path: Path, dest: Path) -> None:
    """Render *src_text* as a DOCX via python-docx."""
    from docx import Document  # python-docx

    document = Document()
    document.add_heading(_title_from_path(src_path), level=1)
    for block in src_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Inside a block, keep line breaks as soft breaks within one paragraph
        para = document.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            if i > 0:
                para.add_run().add_break()
            para.add_run(line)

    document.save(str(dest))


def write_html(src_text: str, src_path: Path, dest: Path) -> None:
    """Wrap *src_text* in a minimal HTML5 document."""
    title = _title_from_path(src_path)
    safe_title = (
        title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )
    paragraphs_html = []
    for block in src_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = (
            block.replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;")
                 .replace("\n", "<br/>\n")
        )
        paragraphs_html.append(f"  <p>{safe}</p>")

    html_doc = (
        "<!DOCTYPE html>\n"
        "<html lang=\"en\">\n"
        "<head>\n"
        "  <meta charset=\"utf-8\"/>\n"
        f"  <title>{safe_title}</title>\n"
        "</head>\n"
        "<body>\n"
        f"  <h1>{safe_title}</h1>\n"
        + "\n".join(paragraphs_html) + "\n"
        "</body>\n"
        "</html>\n"
    )
    dest.write_text(html_doc, encoding="utf-8")


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def diversify(true_dir: Path, *, force: bool, dry_run: bool) -> dict:
    """Rewrite a portion of *.txt files in *true_dir* into PDF / DOCX / HTML.

    Returns a counters dict: {pdf, docx, html, txt, skipped}.
    """
    counters = {"pdf": 0, "docx": 0, "html": 0, "txt": 0, "skipped": 0}

    if not true_dir.exists():
        log(f"ERROR: {true_dir} does not exist — run step 02 first")
        sys.exit(1)

    # Sort for deterministic round-robin assignment
    txt_files = sorted(
        p for p in true_dir.iterdir()
        if p.is_file() and p.suffix.lower() == ".txt"
    )

    if not txt_files:
        log(f"WARN: no .txt files in {true_dir} — already diversified?")
        return counters

    log(f"Processing {len(txt_files)} .txt files in {true_dir}")
    log("Format distribution will be roughly: 25 % each of PDF / DOCX / HTML / TXT")

    for idx, src in enumerate(txt_files):
        target = assign_format(idx)
        if target == "txt":
            counters["txt"] += 1
            continue  # leave it as-is

        dest = src.with_suffix(f".{target}")

        # Skip if target already exists (and we're not forcing)
        if dest.exists() and not force:
            log(f"  SKIP {dest.name} (exists; use --force to overwrite)", indent=1)
            counters["skipped"] += 1
            continue

        if dry_run:
            log(f"  DRY {src.name} → {dest.name}", indent=1)
            counters[target] += 1
            continue

        try:
            text = src.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            log(f"  SKIP {src.name}: cannot read ({exc})", indent=1)
            counters["skipped"] += 1
            continue

        try:
            if target == "pdf":
                write_pdf(text, src, dest)
            elif target == "docx":
                write_docx(text, src, dest)
            elif target == "html":
                write_html(text, src, dest)
            counters[target] += 1
        except Exception as exc:  # noqa: BLE001
            log(f"  ERROR converting {src.name} → {target}: {exc}", indent=1)
            counters["skipped"] += 1
            continue

        # Remove the original .txt so we don't keep duplicate content under
        # two different filenames in the corpus.
        try:
            src.unlink()
        except OSError as exc:
            log(f"  WARN could not remove {src.name}: {exc}", indent=1)

        if (counters["pdf"] + counters["docx"] + counters["html"]) % 10 == 0:
            log(
                f"  … {counters['pdf']} pdf, {counters['docx']} docx, "
                f"{counters['html']} html, {counters['txt']} txt kept",
                indent=1,
            )

    return counters


def print_summary(true_dir: Path, counters: dict) -> None:
    log("")
    log("─" * 50)
    log("TRUE CORPUS — MULTI-FORMAT SUMMARY")
    log("─" * 50)

    # Recount what's on disk now (excludes .gitkeep)
    on_disk: dict[str, int] = {}
    total_bytes = 0
    for p in true_dir.iterdir():
        if not p.is_file() or p.name == ".gitkeep":
            continue
        ext = p.suffix.lower().lstrip(".") or "(no ext)"
        on_disk[ext] = on_disk.get(ext, 0) + 1
        total_bytes += p.stat().st_size

    log(f"  Directory : {true_dir}")
    log(f"  Total size: {human_size(total_bytes)}")
    log("  Files by extension:")
    for ext, count in sorted(on_disk.items(), key=lambda x: -x[1]):
        log(f"      .{ext:<6}: {count}", indent=1)

    log("")
    log("  This run converted:")
    log(f"      → pdf : {counters['pdf']}", indent=1)
    log(f"      → docx: {counters['docx']}", indent=1)
    log(f"      → html: {counters['html']}", indent=1)
    log(f"      kept   : {counters['txt']} as txt", indent=1)
    log(f"      skipped: {counters['skipped']}", indent=1)
    log("─" * 50)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Diversify the K8s true-corpus into PDF / DOCX / HTML / TXT"
    )
    parser.add_argument(
        "--force", action="store_true",
        help="Overwrite existing converted files",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Show what would be converted, but write nothing",
    )
    args = parser.parse_args()

    cfg = load_config()
    true_dir = PROJECT_ROOT / cfg["output"]["true_data_dir"]

    log("=" * 60)
    log("ADV RAG — Step 05: Diversify True-Corpus Formats")
    log("=" * 60)

    counters = diversify(true_dir, force=args.force, dry_run=args.dry_run)
    print_summary(true_dir, counters)


if __name__ == "__main__":
    main()
